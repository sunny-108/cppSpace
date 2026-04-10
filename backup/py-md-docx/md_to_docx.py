#!/usr/bin/env python3
"""
Markdown to DOCX Converter
Converts markdown files to Microsoft Word documents with proper formatting.
Preserves tables, code blocks with light background, and proper typography.
"""

import argparse
import re
import markdown
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# Font settings
BODY_FONT = "Gill Sans MT"
BODY_SIZE = 10
HEADING_FONT = "Gill Sans MT"
CODE_FONT = "Consolas"
CODE_SIZE = 8
CODE_BG_COLOR = "F5F5F5"  # Light gray background for code


def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def create_code_block_table(doc, code_text, language=None):
    """Create a table with light background for code blocks to handle overflow."""
    # Create a single-cell table for the code block
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Set table width to fit page
    table.autofit = False
    for cell in table.columns[0].cells:
        cell.width = Inches(6.5)
    
    cell = table.cell(0, 0)
    set_cell_shading(cell, CODE_BG_COLOR)
    
    # Clear default paragraph and add code
    cell.paragraphs[0].clear()
    
    # Split code into lines to handle overflow
    lines = code_text.strip().split('\n')
    for i, line in enumerate(lines):
        if i == 0:
            para = cell.paragraphs[0]
        else:
            para = cell.add_paragraph()
        
        run = para.add_run(line if line else ' ')  # Use space for empty lines
        run.font.name = CODE_FONT
        run._element.rPr.rFonts.set(qn('w:eastAsia'), CODE_FONT)
        run.font.size = Pt(CODE_SIZE)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Set paragraph formatting for tight line spacing
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
    
    # Add small margin inside the cell
    set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
    
    return table


def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """Set margins for a table cell (in twips, 1440 twips = 1 inch)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    
    for margin_name, margin_value in [('top', top), ('bottom', bottom), 
                                       ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(margin_value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    
    tcPr.append(tcMar)


def setup_styles(doc):
    """Setup document styles."""
    styles = doc.styles
    
    # Normal/Body style
    style = styles['Normal']
    font = style.font
    font.name = BODY_FONT
    font.size = Pt(BODY_SIZE)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
    
    # Heading styles
    heading_sizes = {
        'Heading 1': 18,
        'Heading 2': 16,
        'Heading 3': 14,
        'Heading 4': 12,
        'Heading 5': 11,
        'Heading 6': 10,
    }
    
    for heading_name, size in heading_sizes.items():
        try:
            style = styles[heading_name]
            font = style.font
            font.name = HEADING_FONT
            font.size = Pt(size)
            font.bold = True
            font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)  # Professional blue
            style._element.rPr.rFonts.set(qn('w:eastAsia'), HEADING_FONT)
        except KeyError:
            pass
    
    return doc


def add_inline_code(paragraph, text):
    """Add inline code with light background."""
    run = paragraph.add_run(text)
    run.font.name = CODE_FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), CODE_FONT)
    run.font.size = Pt(CODE_SIZE + 1)
    
    # Add shading for inline code
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), CODE_BG_COLOR)
    run._element.get_or_add_rPr().append(shd)
    
    return run


def process_inline_elements(paragraph, element, doc):
    """Process inline elements like bold, italic, code, links."""
    if isinstance(element, NavigableString):
        text = str(element)
        if text.strip() or text:
            run = paragraph.add_run(text)
            run.font.name = BODY_FONT
            run._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
            run.font.size = Pt(BODY_SIZE)
        return
    
    tag_name = element.name
    
    if tag_name == 'strong' or tag_name == 'b':
        for child in element.children:
            if isinstance(child, NavigableString):
                run = paragraph.add_run(str(child))
                run.bold = True
                run.font.name = BODY_FONT
                run._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
                run.font.size = Pt(BODY_SIZE)
            else:
                process_inline_elements(paragraph, child, doc)
                # Make last run bold
                if paragraph.runs:
                    paragraph.runs[-1].bold = True
    
    elif tag_name == 'em' or tag_name == 'i':
        for child in element.children:
            if isinstance(child, NavigableString):
                run = paragraph.add_run(str(child))
                run.italic = True
                run.font.name = BODY_FONT
                run._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
                run.font.size = Pt(BODY_SIZE)
            else:
                process_inline_elements(paragraph, child, doc)
                if paragraph.runs:
                    paragraph.runs[-1].italic = True
    
    elif tag_name == 'code':
        add_inline_code(paragraph, element.get_text())
    
    elif tag_name == 'a':
        text = element.get_text()
        run = paragraph.add_run(text)
        run.font.name = BODY_FONT
        run._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
        run.font.size = Pt(BODY_SIZE)
        run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)  # Link blue
        run.underline = True
    
    elif tag_name == 'br':
        paragraph.add_run('\n')
    
    else:
        # Process children for unknown tags
        for child in element.children:
            process_inline_elements(paragraph, child, doc)


def process_paragraph(doc, element):
    """Process a paragraph element."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    
    for child in element.children:
        process_inline_elements(para, child, doc)
    
    return para


def process_heading(doc, element, level):
    """Process heading elements."""
    heading_style = f'Heading {level}'
    para = doc.add_paragraph(style=heading_style)
    
    for child in element.children:
        if isinstance(child, NavigableString):
            run = para.add_run(str(child))
        else:
            process_inline_elements(para, child, doc)
    
    return para


def process_list(doc, element, ordered=False, level=0):
    """Process ordered and unordered lists."""
    for li in element.find_all('li', recursive=False):
        # Create bullet/number prefix based on list type and level
        indent = "    " * level
        if ordered:
            # Find index of this li in parent
            idx = list(element.find_all('li', recursive=False)).index(li) + 1
            prefix = f"{idx}. "
        else:
            bullets = ['•', '◦', '▪', '▸']
            prefix = f"{bullets[level % len(bullets)]} "
        
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        para.paragraph_format.space_after = Pt(3)
        
        # Add prefix
        run = para.add_run(prefix)
        run.font.name = BODY_FONT
        run._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
        run.font.size = Pt(BODY_SIZE)
        
        # Process li content
        for child in li.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    run = para.add_run(text)
                    run.font.name = BODY_FONT
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
                    run.font.size = Pt(BODY_SIZE)
            elif child.name in ['ul', 'ol']:
                # Nested list
                process_list(doc, child, ordered=(child.name == 'ol'), level=level + 1)
            elif child.name == 'p':
                for subchild in child.children:
                    process_inline_elements(para, subchild, doc)
            else:
                process_inline_elements(para, child, doc)


def process_table(doc, table_element):
    """Process HTML table and convert to Word table."""
    rows = table_element.find_all('tr')
    if not rows:
        return None
    
    # Count columns from first row
    first_row = rows[0]
    cols = first_row.find_all(['th', 'td'])
    num_cols = len(cols)
    
    if num_cols == 0:
        return None
    
    # Create Word table
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Set table to auto-fit contents
    table.autofit = True
    
    for row_idx, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for col_idx, cell in enumerate(cells):
            if col_idx >= num_cols:
                break
            
            word_cell = table.cell(row_idx, col_idx)
            
            # Clear default paragraph
            word_cell.paragraphs[0].clear()
            
            # Check if header cell
            is_header = cell.name == 'th' or row_idx == 0
            
            # Set header background
            if is_header:
                set_cell_shading(word_cell, 'E7E6E6')
            
            # Get cell text and add to Word cell
            para = word_cell.paragraphs[0]
            
            for child in cell.children:
                if isinstance(child, NavigableString):
                    text = str(child).strip()
                    if text:
                        run = para.add_run(text)
                        run.font.name = BODY_FONT
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
                        run.font.size = Pt(BODY_SIZE - 1)
                        if is_header:
                            run.bold = True
                else:
                    process_inline_elements(para, child, doc)
                    if is_header and para.runs:
                        para.runs[-1].bold = True
            
            # Set cell padding
            set_cell_margins(word_cell, top=50, bottom=50, left=75, right=75)
    
    return table


def process_code_block(doc, element):
    """Process code block (pre > code)."""
    code_element = element.find('code')
    if code_element:
        code_text = code_element.get_text()
        # Try to get language from class
        language = None
        if code_element.get('class'):
            for cls in code_element.get('class'):
                if cls.startswith('language-'):
                    language = cls.replace('language-', '')
                    break
    else:
        code_text = element.get_text()
        language = None
    
    create_code_block_table(doc, code_text, language)
    doc.add_paragraph()  # Add spacing after code block


def process_blockquote(doc, element):
    """Process blockquote elements."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    
    # Add a vertical bar effect using border
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left_border = OxmlElement('w:left')
    left_border.set(qn('w:val'), 'single')
    left_border.set(qn('w:sz'), '24')
    left_border.set(qn('w:space'), '4')
    left_border.set(qn('w:color'), 'CCCCCC')
    pBdr.append(left_border)
    pPr.append(pBdr)
    
    # Process content
    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                run = para.add_run(text)
                run.font.name = BODY_FONT
                run._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
                run.font.size = Pt(BODY_SIZE)
                run.italic = True
        elif child.name == 'p':
            for subchild in child.children:
                process_inline_elements(para, subchild, doc)
        else:
            process_inline_elements(para, child, doc)


def process_hr(doc):
    """Add horizontal rule."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)
    
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def process_element(doc, element):
    """Process a single HTML element and add to document."""
    if isinstance(element, NavigableString):
        text = str(element).strip()
        if text:
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = BODY_FONT
            run._element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
            run.font.size = Pt(BODY_SIZE)
        return
    
    tag_name = element.name
    
    if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(tag_name[1])
        process_heading(doc, element, level)
    
    elif tag_name == 'p':
        process_paragraph(doc, element)
    
    elif tag_name == 'ul':
        process_list(doc, element, ordered=False)
    
    elif tag_name == 'ol':
        process_list(doc, element, ordered=True)
    
    elif tag_name == 'table':
        process_table(doc, element)
        doc.add_paragraph()  # Add spacing after table
    
    elif tag_name == 'pre':
        process_code_block(doc, element)
    
    elif tag_name == 'blockquote':
        process_blockquote(doc, element)
    
    elif tag_name == 'hr':
        process_hr(doc)
    
    elif tag_name in ['div', 'section', 'article']:
        # Process children of container elements
        for child in element.children:
            process_element(doc, child)


def convert_md_to_docx(md_file_path, output_path=None):
    """
    Convert a Markdown file to a Word document.
    
    Args:
        md_file_path: Path to the input markdown file
        output_path: Path for the output docx file (optional)
    
    Returns:
        Path to the created docx file
    """
    # Read markdown file
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert markdown to HTML
    md_extensions = [
        'tables',
        'fenced_code',
        'codehilite',
        'nl2br',
        'sane_lists',
        'smarty',
    ]
    html_content = markdown.markdown(md_content, extensions=md_extensions)
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Create Word document
    doc = Document()
    
    # Setup styles
    setup_styles(doc)
    
    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Process all top-level elements
    for element in soup.children:
        process_element(doc, element)
    
    # Generate output path if not provided
    if output_path is None:
        output_path = md_file_path.rsplit('.', 1)[0] + '.docx'
    
    # Save document
    doc.save(output_path)
    print(f"✅ Successfully converted '{md_file_path}' to '{output_path}'")
    
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description='Convert Markdown files to Microsoft Word documents'
    )
    parser.add_argument(
        'input_file',
        help='Path to the input Markdown file'
    )
    parser.add_argument(
        '-o', '--output',
        help='Path for the output DOCX file (default: same name as input with .docx extension)'
    )
    
    args = parser.parse_args()
    
    try:
        convert_md_to_docx(args.input_file, args.output)
    except FileNotFoundError:
        print(f"❌ Error: File '{args.input_file}' not found")
        exit(1)
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        exit(1)


if __name__ == '__main__':
    main()
